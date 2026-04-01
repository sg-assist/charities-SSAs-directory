#!/usr/bin/env python3
"""Generate a .docx report explaining the UNFPA Partnership Catalyst application."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    if level == 1:
        h.font.size = Pt(22)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table


# ═══════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('UNFPA Partnership Catalyst')
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Application Architecture & Codebase Report')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x00, 0x9E, 0xDB)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Prepared for UNFPA Asia-Pacific\n')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run = meta.add_run('Lee Kuan Yew School of Public Policy — Policy Innovation Lab\n')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run = meta.add_run('April 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════

doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. What the Application Does',
    '3. System Architecture',
    '4. The Knowledge Base',
    '5. The Chat Interface & AI Engine',
    '6. The Ingestion Pipeline',
    '7. Technology Stack',
    '8. Branch & Development Activity',
    '9. Deployment & Infrastructure',
    '10. Purpose of Planned Deliverables (Slides)',
    '11. Team & Attribution',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════

doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'The UNFPA Partnership Catalyst is a web-based research intelligence platform built to support '
    'UNFPA staff in the Asia-Pacific region as they prepare for funder conversations, partnership '
    'pitches, and strategic briefings. It was developed as part of the Lee Kuan Yew School of Public '
    'Policy (LKYSPP) Policy Innovation Lab consulting project for UNFPA, under the supervision of '
    'Professor Mancini.'
)

doc.add_paragraph(
    'The platform combines two capabilities into a single interface:'
)

add_bullet(' A structured, browsable knowledge base of 30+ deep-research documents covering '
           'UNFPA\'s mandate, programme areas, data systems, contested topics, and partnership models.',
           bold_prefix='Knowledge Base: ')
add_bullet(' An AI-powered chat assistant that uses semantic search (RAG) across the knowledge base '
           'and real-time web search to answer complex questions about UNFPA programmes, funders, '
           'and partnership opportunities.',
           bold_prefix='AI Chat: ')

doc.add_paragraph(
    'The application is live at unfpa-lkyspp-otg.vercel.app and is designed for UNFPA programme '
    'officers, country representatives, and partnership teams preparing for meetings with development '
    'finance institutions, philanthropic foundations, and government donors.'
)

# ═══════════════════════════════════════════════════════
# 2. WHAT THE APPLICATION DOES
# ═══════════════════════════════════════════════════════

doc.add_heading('2. What the Application Does', level=1)

doc.add_heading('2.1 The Problem It Solves', level=2)
doc.add_paragraph(
    'UNFPA staff preparing for partnership and funding conversations face a recurring challenge: '
    'institutional knowledge is scattered across hundreds of PDF reports, strategic plans, evaluation '
    'documents, and internal briefings. Finding the right evidence, framing a programme area for a '
    'specific funder, or understanding contested topics requires hours of manual research across '
    'disparate sources.'
)
doc.add_paragraph(
    'The Partnership Catalyst solves this by providing a single interface where staff can:'
)
add_bullet('Browse curated, researcher-written documents on every major UNFPA topic')
add_bullet('Ask natural-language questions and receive evidence-grounded answers')
add_bullet('Get real-time information about funders, markets, and current events via web search')
add_bullet('Generate briefing content, talking points, and funder-specific pitches')

doc.add_heading('2.2 Core User Journeys', level=2)

add_bullet(' A programme officer preparing for a meeting with a Singapore-based family office '
           'asks the chat: "Help me pitch our climate-SRHR programme for a family office interested '
           'in climate adaptation in Southeast Asia." The AI searches the knowledge base for relevant '
           'programme documents and searches the web for the funder\'s recent interests, then generates '
           'a tailored pitch.',
           bold_prefix='Funder Pitch Preparation: ')

add_bullet(' A country representative needs to brief a visiting delegation. They browse the knowledge '
           'base for the relevant programme documents, read the contested areas section to anticipate '
           'difficult questions, and use the chat to generate talking points.',
           bold_prefix='Briefing Preparation: ')

add_bullet(' A partnership team wants to understand how UNFPA\'s midwifery programme could be framed '
           'within a blended finance structure. The chat searches both the knowledge base (for UNFPA '
           'programme details and PPP models) and the web (for current blended finance trends).',
           bold_prefix='Partnership Structuring: ')

# ═══════════════════════════════════════════════════════
# 3. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════

doc.add_heading('3. System Architecture', level=1)

doc.add_paragraph(
    'The application follows a modern full-stack architecture with server-side rendering, '
    'API routes, and external service integrations.'
)

doc.add_heading('3.1 High-Level Architecture', level=2)

add_table(
    ['Layer', 'Technology', 'Purpose'],
    [
        ['Frontend', 'Next.js 16 + React 19', 'Server-rendered pages, interactive chat UI'],
        ['API Layer', 'Next.js API Routes', 'Chat endpoint, admin endpoints, quota checks'],
        ['AI Engine', 'Claude Sonnet 4 (Anthropic)', 'Agentic chat with tool use and extended thinking'],
        ['Embeddings', 'OpenAI text-embedding-3-small', 'Vector embeddings for semantic search (1536 dims)'],
        ['Database', 'PostgreSQL + pgvector', 'Document storage, vector search, PDF tracking'],
        ['Rate Limiting', 'Upstash Redis', 'Daily query quotas (20/day per user)'],
        ['Hosting', 'Vercel', 'Serverless deployment with edge functions'],
        ['Web Search', 'Anthropic web_search tool', 'Real-time information retrieval'],
    ]
)

doc.add_heading('3.2 Repository Structure', level=2)

doc.add_paragraph(
    'The repository is organised into two main sections:'
)

add_bullet(' Contains all 30+ researcher-written markdown documents, '
           'organised by thematic block (Orientation, Programme Work, Data & Evidence, '
           'Contested Areas, PMNCH, Resilience & Partnerships). Also contains the PDF source '
           'library with 100+ reports.',
           bold_prefix='docs/ — Knowledge Base Content: ')

add_bullet(' The full Next.js web application including the chat interface, '
           'knowledge base browser, API routes, ingestion scripts, database schema, and all '
           'supporting services.',
           bold_prefix='next-app/ — Web Application: ')

# ═══════════════════════════════════════════════════════
# 4. THE KNOWLEDGE BASE
# ═══════════════════════════════════════════════════════

doc.add_heading('4. The Knowledge Base', level=1)

doc.add_paragraph(
    'The knowledge base is the intellectual core of the application. It is a two-layer system '
    'designed to provide both synthesised analysis and primary source depth.'
)

doc.add_heading('4.1 Layer 2 — Structured Research Documents (Built First)', level=2)
doc.add_paragraph(
    '30 researcher-written documents, each 3,000-14,000 words, covering UNFPA\'s mandate, programmes, '
    'evidence base, and politically sensitive topics. These are not summaries of existing reports — '
    'they are original analytical documents written to answer specific questions that practitioners '
    'and decision-makers actually ask.'
)

doc.add_paragraph('The documents are organised into six thematic blocks:')

add_table(
    ['Block', 'Code', 'Documents', 'Description'],
    [
        ['Orientation', 'O', '8', 'What UNFPA and PMNCH are, how they work, key terminology, mandate foundations'],
        ['Programme Work', 'W', '10', 'Deep dives into maternal health, family planning, GBV, fistula, midwifery, FGM, child marriage, adolescent SRH, contraceptive procurement'],
        ['Data & Evidence', 'D', '4', 'SOWP reports, census/CRVS support, demographic dividend, results reporting methodology'],
        ['Contested Areas', 'C', '5', 'US defunding episodes, abortion mandate, CSE controversy, disputed results, China programme'],
        ['PMNCH', 'PMNCH', '5', 'Partnership for Maternal, Newborn & Child Health — mandate, accountability, research, relationship to UNFPA'],
        ['Resilience', 'R', '4', 'PPP models, climate-SRHR nexus, Singapore financial ecosystem, community resilience'],
    ]
)

doc.add_heading('4.2 Layer 1 — Ingested PDF Reports (Supplementary)', level=2)
doc.add_paragraph(
    'A library of 100+ PDF source documents downloaded from UNFPA, PMNCH, WHO, and partner organisations. '
    'These include annual reports (1997-2024), strategic plans, maternal health research, evaluation reports, '
    'and technical guidance. They are ingested with SHA-256 change detection, parsed, chunked, and embedded '
    'for semantic search. Layer 1 extends the depth of Layer 2 by providing primary source detail.'
)

doc.add_heading('4.3 What Makes This Knowledge Base Different', level=2)
add_bullet('Each document explicitly distinguishes between what UNFPA reports, '
           'what independent evaluations find, and where the evidence is genuinely uncertain.',
           bold_prefix='Honest Assessment: ')
add_bullet('Documents cover US defunding, the abortion mandate, CSE controversy, '
           'disputed results, and the China programme — topics most institutional resources avoid.',
           bold_prefix='Contested Topics Included: ')
add_bullet('Content is calibrated for two distinct audiences: practitioners (field staff, '
           'programme officers) and decision-makers (senior managers, donors, board members).',
           bold_prefix='Audience-Aware: ')
add_bullet('Documents cross-reference each other with document codes, '
           'creating an interconnected analytical framework.',
           bold_prefix='Cross-Referenced: ')

# ═══════════════════════════════════════════════════════
# 5. THE CHAT INTERFACE & AI ENGINE
# ═══════════════════════════════════════════════════════

doc.add_heading('5. The Chat Interface & AI Engine', level=1)

doc.add_heading('5.1 How the Chat Works', level=2)
doc.add_paragraph(
    'The chat interface is powered by an agentic AI system that combines three capabilities:'
)

add_bullet(' Claude searches the embedded knowledge base '
           'using vector similarity (pgvector cosine distance) to find the most relevant document '
           'chunks for any query. This grounds responses in the curated research.',
           bold_prefix='Semantic Knowledge Base Search: ')

add_bullet(' Claude can search the live web using Anthropic\'s '
           'built-in web search tool to find current information about funders, market conditions, '
           'recent policy changes, and news relevant to the user\'s query.',
           bold_prefix='Real-Time Web Search: ')

add_bullet(' Claude uses a dedicated reasoning phase (8,000-token '
           'budget) to think through complex questions before responding, ensuring nuanced answers '
           'for multi-faceted partnership and funding questions.',
           bold_prefix='Extended Thinking: ')

doc.add_heading('5.2 The Agentic Loop', level=2)
doc.add_paragraph(
    'Unlike a simple question-answer system, the chat operates as an agentic loop:'
)
doc.add_paragraph(
    '1. The user submits a question.\n'
    '2. Claude analyses the question with extended thinking.\n'
    '3. Claude decides which tools to use — knowledge base search, web search, or both.\n'
    '4. Tool results are returned to Claude, which may decide to search again with refined queries.\n'
    '5. This loop continues for up to 6 rounds until Claude has sufficient information.\n'
    '6. Claude generates a comprehensive, source-attributed response.\n'
    '7. If the response is truncated, wave generation automatically continues it (up to 4 waves).'
)

doc.add_heading('5.3 Streaming & User Experience', level=2)
doc.add_paragraph(
    'Responses are streamed to the user via Server-Sent Events (SSE) with real-time status indicators '
    'showing the current phase: "Thinking...", "Searching knowledge base...", "Searching the web...", '
    '"Writing...". This provides transparency into the AI\'s reasoning process.'
)

doc.add_heading('5.4 Starter Prompts', level=2)
doc.add_paragraph(
    'The interface provides six pre-built starter prompts designed for common UNFPA partnership scenarios:'
)
add_bullet('Pitch a climate adaptation programme to a Singapore family office')
add_bullet('Generate a briefing note on climate-SRHR for philanthropic partners')
add_bullet('Prepare talking points for a development finance institution meeting')
add_bullet('Match UNFPA projects to potential funders in Southeast Asia')
add_bullet('Frame SRHR within climate funding mechanisms')
add_bullet('Compare blended finance, impact bonds, and South-South cooperation models')

doc.add_heading('5.5 Source Attribution', level=2)
doc.add_paragraph(
    'Every response includes clickable source links to the knowledge base documents used. '
    'This allows users to verify claims, read the full context, and build confidence in the '
    'AI\'s outputs. Sources are tracked throughout the agentic loop and displayed alongside '
    'the final response.'
)

# ═══════════════════════════════════════════════════════
# 6. THE INGESTION PIPELINE
# ═══════════════════════════════════════════════════════

doc.add_heading('6. The Ingestion Pipeline', level=1)

doc.add_paragraph(
    'The system uses a sophisticated multi-stage pipeline to convert raw documents into searchable '
    'vector embeddings.'
)

doc.add_heading('6.1 Pipeline Stages', level=2)

add_table(
    ['Stage', 'Process', 'Details'],
    [
        ['1. Discovery', 'File scanning', 'Reads .md and .pdf files from the docs/ directory'],
        ['2. Parsing', 'Content extraction', 'Markdown: frontmatter + body. PDFs: text extraction with ligature fixing, header/footer removal'],
        ['3. Change Detection', 'SHA-256 hashing', 'Computes content hash; skips unchanged files to avoid redundant processing'],
        ['4. Chunking', 'Intelligent splitting', 'Splits on heading boundaries (markdown) or paragraphs (PDFs); ~800-1200 words per chunk with 100-word overlap'],
        ['5. Embedding', 'Vector generation', 'OpenAI text-embedding-3-small generates 1536-dimension vectors for each chunk'],
        ['6. Storage', 'Database insertion', 'Documents, chunks, and embeddings stored in PostgreSQL with pgvector'],
    ]
)

doc.add_heading('6.2 Smart Change Detection', level=2)
doc.add_paragraph(
    'Both markdown and PDF ingestion use SHA-256 content hashing to detect changes. When a document '
    'is updated, only that document is re-chunked and re-embedded. This makes the ingestion process '
    'efficient and idempotent — it can be run repeatedly without wasting resources on unchanged content.'
)

# ═══════════════════════════════════════════════════════
# 7. TECHNOLOGY STACK
# ═══════════════════════════════════════════════════════

doc.add_heading('7. Technology Stack', level=1)

add_table(
    ['Category', 'Technology', 'Version / Details'],
    [
        ['Framework', 'Next.js', 'v16 with App Router'],
        ['UI Library', 'React', 'v19'],
        ['Language', 'TypeScript', 'v5'],
        ['Database', 'PostgreSQL + pgvector', 'Supabase-hosted, 1536-dim vectors'],
        ['ORM', 'Prisma', 'v7'],
        ['LLM', 'Claude Sonnet 4', 'Anthropic API with tool use'],
        ['Embeddings', 'OpenAI text-embedding-3-small', '1536 dimensions'],
        ['Web Search', 'Anthropic web_search_20250305', 'Server-side tool'],
        ['Rate Limiting', 'Upstash Redis', '20 queries/day'],
        ['Styling', 'Tailwind CSS', 'v4 with typography plugin'],
        ['PDF Parsing', 'pdf-parse', 'Text extraction with cleaning'],
        ['Markdown', 'marked + react-markdown', 'Server & client rendering'],
        ['Hosting', 'Vercel', 'Serverless with edge functions'],
        ['Icons', 'Lucide React', 'Open-source icon library'],
    ]
)

# ═══════════════════════════════════════════════════════
# 8. BRANCH & DEVELOPMENT ACTIVITY
# ═══════════════════════════════════════════════════════

doc.add_heading('8. Branch & Development Activity', level=1)

doc.add_paragraph(
    'The repository has two active branches reflecting the development history:'
)

doc.add_heading('8.1 Main Branch', level=2)
doc.add_paragraph(
    'The stable production branch. Contains the foundational application with the knowledge base browser, '
    'chat interface, rate limiting, and the initial product positioning as an UNFPA partnership tool. '
    'Includes the PDF library manifests and bulk download tooling.'
)

doc.add_heading('8.2 Merged Feature Branches', level=2)

add_bullet(' Reframed the product from an internal LKYSPP research tool '
           'to a UNFPA Asia-Pacific-facing partnership preparation tool. Updated the system prompt, '
           'added partnership-focused starter prompts, rebranded to "UNFPA Partnership Catalyst", and '
           'updated attribution.',
           bold_prefix='PR #1 — Product Repositioning: ')

add_bullet(' Major upgrade to the chat system. Rewrote the '
           'Chat API with an agentic tool-use loop, added Anthropic\'s web search capability, enabled '
           'extended thinking with 8K token budget, implemented wave generation for long responses, '
           'and added SSE-based streaming with progressive status indicators. Also added SHA-256 '
           'content hashing for the knowledge base ingestion.',
           bold_prefix='PR #2 — Agentic Chat & Knowledge Base Ingestion: ')

# ═══════════════════════════════════════════════════════
# 9. DEPLOYMENT & INFRASTRUCTURE
# ═══════════════════════════════════════════════════════

doc.add_heading('9. Deployment & Infrastructure', level=1)

doc.add_paragraph(
    'The application is designed for zero-ops deployment on Vercel with managed database services.'
)

add_bullet(' Deployed on Vercel with the root directory set to next-app/. '
           'Knowledge base pages are statically rendered at build time (no database required for browsing). '
           'The chat API runs as a serverless function.',
           bold_prefix='Hosting: ')

add_bullet(' PostgreSQL with pgvector extension, hosted on Supabase free tier. '
           'Stores documents, chunks, and vector embeddings.',
           bold_prefix='Database: ')

add_bullet(' Upstash Redis for daily query rate limiting (20 queries/day). '
           'Gracefully degrades if Redis is unavailable.',
           bold_prefix='Rate Limiting: ')

add_bullet(' Requires four environment variables: DATABASE_URL (PostgreSQL), '
           'ANTHROPIC_API_KEY (chat), OPENAI_API_KEY (embeddings), and optionally UPSTASH_REDIS_REST_URL '
           'and UPSTASH_REDIS_REST_TOKEN (rate limiting).',
           bold_prefix='Configuration: ')

add_bullet(' The code is MIT-licensed. The knowledge base documents are CC BY 4.0.',
           bold_prefix='Licensing: ')

# ═══════════════════════════════════════════════════════
# 10. PURPOSE OF PLANNED DELIVERABLES (SLIDES)
# ═══════════════════════════════════════════════════════

doc.add_heading('10. Purpose of Planned Deliverables (Slides)', level=1)

doc.add_paragraph(
    'The accompanying PowerPoint presentation is designed as a client briefing deck for UNFPA. '
    'Its purpose is to:'
)

add_bullet('Introduce the Partnership Catalyst to UNFPA stakeholders who have not yet seen the tool')
add_bullet('Explain what the knowledge base contains and why it was structured this way')
add_bullet('Demonstrate the AI chat capability and its practical value for partnership preparation')
add_bullet('Provide a clear overview of the technical architecture for IT and procurement teams')
add_bullet('Outline next steps and potential expansion paths')

doc.add_paragraph(
    'The slides are intentionally designed to be editable — structured with clear sections, '
    'minimal text, and placeholders where screenshots or live demo references can be added '
    'before the client presentation.'
)

# ═══════════════════════════════════════════════════════
# 11. TEAM & ATTRIBUTION
# ═══════════════════════════════════════════════════════

doc.add_heading('11. Team & Attribution', level=1)

doc.add_heading('11.1 LKYSPP Student Group', level=2)

add_table(
    ['Name', 'Role'],
    [
        ['Rani Opula Rajan', 'Group Member'],
        ['Prachi Sharma', 'Group Member'],
        ['Abhishek Tiwari', 'Group Member'],
        ['Preeti Patil', 'Group Member'],
    ]
)

doc.add_heading('11.2 Application', level=2)
doc.add_paragraph(
    'Application designed and built by Haojun See (MPP 2021) to support the group\'s research '
    'and consulting work for UNFPA.'
)

doc.add_heading('11.3 Academic Context', level=2)
doc.add_paragraph(
    'This project was conducted through Professor Mancini\'s Policy Innovation Lab course at the '
    'Lee Kuan Yew School of Public Policy, National University of Singapore. The client is UNFPA '
    '(United Nations Population Fund), Challenge B: PPP for Climate and Humanitarian Resilience.'
)

# ═══════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════

out_path = os.path.join(os.path.dirname(__file__), '..', 'docs', 'deliverables',
                        'UNFPA-Partnership-Catalyst-Report.docx')
doc.save(out_path)
print(f'Saved: {os.path.abspath(out_path)}')
